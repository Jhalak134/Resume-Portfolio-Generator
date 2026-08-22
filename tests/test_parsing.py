import os
import sys
import tempfile
import unittest

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import docx
import pypdf

from parsing import extract_photo_from_bytes, extract_text_from_bytes, read_resume


class TestTxtExtraction(unittest.TestCase):
    def test_valid_txt(self):
        text = extract_text_from_bytes(
            "resume.txt",
            b"John Doe\nSoftware Engineer\nSkills: Python, Flask\n" + b"x" * 50,
        )
        self.assertIn("John Doe", text)
        self.assertIn("Skills: Python, Flask", text)

    def test_empty_txt_rejected(self):
        with self.assertRaises(ValueError):
            extract_text_from_bytes("resume.txt", b"   \n\n  ")

    def test_too_short_txt_rejected(self):
        with self.assertRaises(ValueError):
            extract_text_from_bytes("resume.txt", b"hi")


class TestPdfExtraction(unittest.TestCase):
    def test_blank_pdf_yields_no_usable_text(self):
        writer = pypdf.PdfWriter()
        writer.add_blank_page(width=612, height=792)
        with tempfile.SpooledTemporaryFile() as buf:
            writer.write(buf)
            buf.seek(0)
            data = buf.read()

        # A blank page has no text, so validation should reject it.
        with self.assertRaises(ValueError):
            extract_text_from_bytes("resume.pdf", data)

    def test_corrupt_pdf_raises(self):
        with self.assertRaises(Exception):
            extract_text_from_bytes("resume.pdf", b"not a real pdf")


class TestDocxExtraction(unittest.TestCase):
    def test_valid_docx(self):
        doc = docx.Document()
        doc.add_heading("Jane Doe", level=1)
        doc.add_paragraph("AI Engineer")
        doc.add_paragraph("Experience in Python and PyTorch, building resumes for testing.")

        with tempfile.SpooledTemporaryFile() as buf:
            doc.save(buf)
            buf.seek(0)
            data = buf.read()

        text = extract_text_from_bytes("resume.docx", data)
        self.assertIn("Jane Doe", text)
        self.assertIn("AI Engineer", text)

    def test_corrupt_docx_raises(self):
        with self.assertRaises(Exception):
            extract_text_from_bytes("resume.docx", b"not a real docx")


class TestUnsupportedFormat(unittest.TestCase):
    def test_unsupported_extension_rejected(self):
        with self.assertRaises(ValueError) as ctx:
            extract_text_from_bytes("resume.png", b"fake image data")
        self.assertIn("Unsupported file type", str(ctx.exception))


class TestPhotoExtraction(unittest.TestCase):
    def test_txt_has_no_photo(self):
        photo = extract_photo_from_bytes("resume.txt", b"just text")
        self.assertEqual(photo, "")

    def test_docx_without_image_returns_empty(self):
        doc = docx.Document()
        doc.add_paragraph("No photo here")
        with tempfile.SpooledTemporaryFile() as buf:
            doc.save(buf)
            buf.seek(0)
            data = buf.read()
        photo = extract_photo_from_bytes("resume.docx", data)
        self.assertEqual(photo, "")


class TestReadResume(unittest.TestCase):
    def test_reads_sample_resume(self):
        repo_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        sample_path = os.path.join(repo_root, "sample_resume.txt")
        if os.path.exists(sample_path):
            text = read_resume(sample_path)
            self.assertIn("John Malik", text)

    def test_missing_file_raises(self):
        with self.assertRaises(ValueError):
            read_resume("this_file_does_not_exist.txt")


if __name__ == "__main__":
    unittest.main()