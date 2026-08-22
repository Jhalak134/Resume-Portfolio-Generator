"""DOCX resume extraction: plain text and a best-effort embedded photo."""

import base64
import io


def extract_text(data: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(data))
    return "\n".join(p.text for p in document.paragraphs)


def extract_photo(data: bytes) -> str:
    """Return a base64 data URI for the first embedded image found, or ""."""
    import docx

    try:
        document = docx.Document(io.BytesIO(data))
        for rel in document.part.rels.values():
            if "image" in rel.reltype:
                image_part = rel.target_part
                mime = image_part.content_type or "image/png"
                b64 = base64.b64encode(image_part.blob).decode("ascii")
                return f"data:{mime};base64,{b64}"
    except Exception:
        pass
    return ""
